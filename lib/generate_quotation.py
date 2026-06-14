import sys
import json
import os
from datetime import date
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
TNR = f'<w:rFonts xmlns:w="{NS}" w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>'

def xe(s):
    return str(s).replace("&","&amp;").replace("<","&lt;").replace(">","&gt;").replace('"',"&quot;")

def make_product_row(sno, name, desc, qty, price_excl, total):
    instr = xe(r' =PRODUCT(LEFT) \# "#,##0.00" ')
    return etree.fromstring(f'''<w:tr xmlns:w="{NS}">
  <w:trPr><w:trHeight w:val="199"/></w:trPr>
  <w:tc><w:tcPr><w:tcW w:w="918"  w:type="dxa"/></w:tcPr><w:p><w:pPr><w:rPr>{TNR}</w:rPr></w:pPr><w:r><w:rPr>{TNR}</w:rPr><w:t>{xe(sno)}</w:t></w:r></w:p></w:tc>
  <w:tc><w:tcPr><w:tcW w:w="2970" w:type="dxa"/></w:tcPr><w:p><w:pPr><w:rPr>{TNR}</w:rPr></w:pPr><w:r><w:rPr>{TNR}</w:rPr><w:t>{xe(name)}</w:t></w:r></w:p></w:tc>
  <w:tc><w:tcPr><w:tcW w:w="2689" w:type="dxa"/></w:tcPr><w:p><w:pPr><w:rPr>{TNR}</w:rPr></w:pPr><w:r><w:rPr>{TNR}</w:rPr><w:t>{xe(desc)}</w:t></w:r></w:p></w:tc>
  <w:tc><w:tcPr><w:tcW w:w="731"  w:type="dxa"/></w:tcPr><w:p><w:pPr><w:rPr>{TNR}</w:rPr></w:pPr><w:r><w:rPr>{TNR}</w:rPr><w:t>{xe(str(qty))}</w:t></w:r></w:p></w:tc>
  <w:tc><w:tcPr><w:tcW w:w="1530" w:type="dxa"/></w:tcPr><w:p><w:pPr><w:rPr>{TNR}</w:rPr></w:pPr><w:r><w:rPr>{TNR}</w:rPr><w:t>{price_excl:.2f}</w:t></w:r></w:p></w:tc>
  <w:tc><w:tcPr><w:tcW w:w="1735" w:type="dxa"/></w:tcPr>
    <w:p><w:pPr><w:rPr>{TNR}<w:b/><w:bCs/></w:rPr></w:pPr>
      <w:fldSimple w:instr="{instr}"><w:r><w:rPr>{TNR}<w:b/><w:bCs/></w:rPr><w:t>{total:,.2f}</w:t></w:r></w:fldSimple>
    </w:p></w:tc>
</w:tr>''')

def make_summary_row(label, instr, bold=False, top_border=False):
    b   = f'<w:b xmlns:w="{NS}"/><w:bCs xmlns:w="{NS}"/>' if bold else ''
    bdr = (f'<w:tcBorders xmlns:w="{NS}"><w:top w:val="single" w:sz="6" w:space="0" w:color="000000"/>'
           f'</w:tcBorders>') if top_border else ''
    return etree.fromstring(f'''<w:tr xmlns:w="{NS}">
  <w:trPr><w:trHeight w:val="260"/></w:trPr>
  <w:tc>
    <w:tcPr><w:tcW w:w="8037" w:type="dxa"/><w:gridSpan w:val="5"/>{bdr}</w:tcPr>
    <w:p>
      <w:pPr><w:jc w:val="right"/><w:rPr>{TNR}{b}<w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr></w:pPr>
      <w:r><w:rPr>{TNR}{b}<w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr>
        <w:t xml:space="preserve">{xe(label)}</w:t></w:r>
    </w:p>
  </w:tc>
  <w:tc>
    <w:tcPr><w:tcW w:w="1735" w:type="dxa"/>{bdr}</w:tcPr>
    <w:p>
      <w:pPr><w:rPr>{TNR}{b}<w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr></w:pPr>
      <w:fldSimple w:instr="{xe(instr)}">
        <w:r><w:rPr>{TNR}{b}<w:sz w:val="24"/><w:szCs w:val="24"/></w:rPr><w:t>0.00</w:t></w:r>
      </w:fldSimple>
    </w:p>
  </w:tc>
</w:tr>''')

def main():
    try:
        # Read JSON from stdin
        data = json.loads(sys.stdin.read())
        template_path = data["template_path"]
        output_path = data["output_path"]
        items = data["items"]

        if not os.path.exists(template_path):
            print(f"Error: Template not found at {template_path}", file=sys.stderr)
            sys.exit(1)

        doc = Document(template_path)

        # Update date to today automatically
        today = date.today().strftime("%d/%m/%Y")
        date_cell = doc.tables[1].rows[0].cells[0]
        for para in date_cell.paragraphs:
            runs = para.runs
            if runs and "Date:" in runs[0].text:
                runs[0].text = f"Date: {today}"
                for run in runs[1:]:
                    run.text = ""
                break

        tbl = doc.tables[2]
        tbl_el = tbl._tbl

        # Remove all rows except header
        for tr in list(tbl_el.findall(qn('w:tr')))[1:]:
            tbl_el.remove(tr)

        # Add product rows
        for idx, item in enumerate(items):
            sno = f"{idx + 1}."
            name = item["name"]
            desc = item.get("description", "")
            qty = int(item["qty"])
            price_incl = float(item["priceWithGst"])
            price_excl = price_incl / 1.18
            tbl_el.append(make_product_row(sno, name, desc, qty, price_excl, qty * price_excl))

        # Add summary rows (dynamic range)
        n = len(items)
        f_range = f"F2:F{1 + n}"

        tbl_el.append(make_summary_row("Sub Total (Without GST)  \u20b9",
            r' =SUM(ABOVE) \# "#,##0.00" ', bold=True, top_border=True))
        tbl_el.append(make_summary_row("CGST @ 9%  \u20b9",
            f' =SUM({f_range})*0.09 \\# "#,##0.00" '))
        tbl_el.append(make_summary_row("SGST @ 9%  \u20b9",
            f' =SUM({f_range})*0.09 \\# "#,##0.00" '))
        tbl_el.append(make_summary_row("Grand Total  (Incl. 18% GST)  \u20b9",
            f' =SUM({f_range})*1.18 \\# "#,##0.00" ', bold=True))

        # Remove old hardcoded CGST paragraph
        body = doc.element.body
        for para in body.findall(qn('w:p')):
            text = ''.join(t.text or '' for t in para.iter(qn('w:t')))
            if 'CGST' in text and ('1315' in text or '\u2013' in text):
                body.remove(para)
                break

        doc.save(output_path)
        print("Success")
        sys.exit(0)

    except Exception as e:
        import traceback
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)

if __name__ == "__main__":
    main()
